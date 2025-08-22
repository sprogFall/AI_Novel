import os
import requests
import json
import time
import re
import queue
import threading
import uuid
import shutil
from datetime import datetime
from io import BytesIO
from flask import Flask, render_template_string, request, jsonify, send_file
from markdown import markdown
# NOTE: This application requires the 'python-docx' and 'beautifulsoup4' libraries.
# Install with: pip install python-docx beautifulsoup4
from docx import Document
from bs4 import BeautifulSoup

# --- LLM API Configuration ---
# 请根据你的环境配置这些设置
<<<<<<< Updated upstream
VLLM_SERVER_HOST = ""
VLLM_SERVER_PORT = ""
OPENAI_API_ENDPOINT_PATH = "/v1/chat/completions"
TARGET_API_URL = f"http://{VLLM_SERVER_HOST}:{VLLM_SERVER_PORT}{OPENAI_API_ENDPOINT_PATH}"
SERVED_MODEL_IDENTIFIER = ''
=======
VLLM_SERVER_HOST = "ip"
VLLM_SERVER_PORT = "port"
OPENAI_API_ENDPOINT_PATH = "/v1/chat/completions"
TARGET_API_URL = f"http://{VLLM_SERVER_HOST}:{VLLM_SERVER_PORT}{OPENAI_API_ENDPOINT_PATH}"
SERVED_MODEL_IDENTIFIER = 'models'
>>>>>>> Stashed changes
API_HEADERS = {
    "Content-Type": "application/json",
    # "Authorization": "Bearer YOUR_API_KEY" # 如果需要，请添加你的API密钥
}

# --- 所有生成作品的基础目录 ---
WORKS_LIBRARY_PATH = "works_library"
WORLDVIEW_DIR_NAME = "_worldview"
WORLDVIEW_FILE_NAME = "worldview.md"
MIND_LIBRARY_DIR_NAME = "_mind_library"
MIND_LIBRARY_FILE_NAME = "mind_library.json"
TOKEN_STATS_FILE = "token_stats.json"

OUTLINE_PROMPT_TEMPLATES = {
    "course": {
        "default": """
        请你扮演一位资深的课程设计专家。
        我需要一个关于 "{topic}" 的课程大纲。
        请遵循以下结构和要求：
        1.  在第一行，请以 "标题：[这里是课程标题]" 的格式，为这个课程给出一个简洁、明确的标题。
        2.  从第二行开始，生成详细的课程大纲。
        3.  大纲应包含 {chapters} 个主要章节。
        4.  每个章节下应包含 {sections} 个小节。
        5.  每个小节都需要有一个明确、具体的标题。
        请以Markdown格式输出完整的大纲，章节使用'###'，小节使用'1.'、'2.'等数字列表。
        """
    },
    "novel": {
        "default": """
        请你扮演一位经验丰富的小说编辑。
        我正在构思一部关于 "{topic}" 的小说，请帮我生成一份故事大纲。
        请遵循以下结构和要求：
        1.  在第一行，请以 "书名：[这里是小说标题]" 的格式，为这部小说想一个引人入胜的标题。
        2.  从第二行开始，生成详细的故事大纲。
        3.  小说分为 {chapters} 个主要章节。
        4.  每个章节标题应能概括该章节的核心事件或情感转折。
        请以Markdown格式输出完整的大纲，章节使用'###'。
        """
    }
}

CONTENT_PROMPT_TEMPLATES = {
    "course": """
    请你扮演一位资深的课程研发专家和讲师。
    你的任务是为 "{topic}" 这个课程下的一个小节撰写一份详细、专业且通俗易懂的教学讲义。

    **小节标题**: "{section_title}"

    **内容要求**:
    1.  **格式**: 严格使用 Markdown 格式。
    2.  **结构**: 内容应包含引言、核心概念讲解、关键技术点、代码或伪代码示例（如果适用）、实际案例分析、以及小结。
    3.  **深度**: 内容需要有足够的深度，覆盖该知识点的核心原理和关键细节。
    4.  **专业性**: 术语使用准确，讲解系统化。
    5.  **篇幅**: 请确保内容详实，提供丰富的示例和深入的解释。

    请开始为 "{section_title}" 这一节撰写讲义内容。
    """,
    "novel": """
    请你扮演一位才华横溢、文笔细腻的小说家。
    你的任务是完全按照给定的“本章剧本”，并融入丰富的文学技巧，创作出小说《{topic}》中指定章节的完整内容。

    **核心任务**: 续写标题为“{section_title}”的章节。

    **背景信息 (供参考)**:
    1.  **宏大世界观 (摘要)**:
        ```
        {worldview_context}
        ```
    2.  **上一章节回顾 (精简摘要)**:
        ```
        {previous_chapter_context}
        ```

    **创作指令 (必须严格遵守)**:
    1.  **本章剧本 (来自思维库)**:
        ```
        {chapter_plan}
        ```
    2.  **绝对遵循剧本**: 你的创作必须**严格且完全**基于“本章剧本”进行。剧本中提到的剧情点、人物互动、场景和要引入的新信息都必须在正文中得到体现。**不要偏离剧本，不要即兴发挥**。
    
    **文笔与技巧要求 (非常重要)**:
    1.  **展示而非告知 (Show, Don't Tell)**: 不要直接说“他很害怕”，而是通过描写他“手心出汗、呼吸急促、眼神躲闪”来展现他的害怕。用动作、场景、感官细节来传达情感和信息。
    2.  **丰富感官描写**: 充分描写角色所看到的、听到的、闻到的、尝到的和感觉到的东西，让读者身临其境。
    3.  **避免重复**: 尤其是对于核心设定（如主角的能力），请使用多样化的词汇和比喻进行描述，避免在每一章都使用完全相同的形容词。
    4.  **个性化对话**: 让每个角色的说话方式都符合其在世界观中的设定。一个神秘的导师和一个鲁莽的战士，他们的用词、语气和节奏应该完全不同。对话不仅要推动情节，更要塑造人物。
    5.  **强化内心独白**: 深入主角的内心世界，通过他的思考、挣扎和自我怀疑，来展现人物的成长和弧光，让角色更加立体。

    **最终要求**:
    - 直接创作小说正文，专注于场景、动作、对话和心理描写。不要在正文中包含任何元注释、分析或标题。
    - 保持生动的文笔，并确保本章内容详尽，字数在 {word_count} 字左右。

    现在，请严格依据“本章剧本”和上述所有技巧要求，开始创作《{topic}》中标题为“{section_title}”的章节。
    """,
    "generate_chapter_script": """
    请你扮演一位心思缜密的小说剧情规划师。
    你的任务是基于已有的宏大世界观和刚刚结束的章节，为即将开始的新章节制定一个详细的执行剧本。

    **1. 宏大世界观 (世界规则与核心设定)**:
    ```
    {worldview}
    ```

    **2. 上一章剧情回顾 (精简摘要)**:
    ```
    {previous_chapter_summary}
    ```

    **3. 本章标题**: "{section_title}"

    **你的任务是为 "{section_title}" 这一章，规划出清晰、可执行的写作蓝图。请严格按照以下JSON格式输出**:
    ```json
    {{
      "chapter_title": "{section_title}",
      "status": "pending",
      "summary": "用2-3句话概括本章的核心故事脉络。",
      "key_events": [
        "列出本章必须发生的第1个关键事件。",
        "列出本章必须发生的第2个关键事件。",
        "列出本章必须发生的第3个关键事件。"
      ],
      "登場人物": ["列出本章的主要出场人物"]
    }}
    ```
    请确保输出是**一个完整且格式正确**的JSON对象，不要包含任何额外的解释或Markdown标记。
    """,
    "summarizer": """
    请将以下小说章节内容浓缩成一段150字以内的精简摘要，用于帮助AI理解上下文。摘要需要清晰地概括出以下几点：
    1.  本章的核心事件是什么？
    2.  主要人物在本章结束时的状态（物理、情感）是怎样的？
    3.  是否留下了悬念或为下一章铺垫了哪些线索？

    **章节内容**:
    ---
    {text_to_summarize}
    ---
    
    请输出精简摘要。
    """,
    "initial_worldview": """
    请你扮演一位世界级的小说策划师和设定师。
    你的任务是根据提供的小说主题和高级大纲，构建一个全面且详细的“世界观圣经” (Worldview Bible)。这份文档是后续所有章节创作的基石，必须确保设定的严谨性和一致性。

    **小说主题**: "{topic}"

    **故事大纲**:
    ```
    {outline}
    ```

    **请在世界观圣经中包含以下核心部分，并使用Markdown格式化**:

    ### 核心概念
    用一到两句话总结整个故事的核心冲突和主题。

    ### 主要人物
    为每位主要人物（至少2-3位）建立档案，每个档案使用 '---' 分隔。档案包括：
    - **姓名**:
    - **身份/背景**:
    - **外貌特征**:
    - **性格特点**: (例如：勇敢但鲁莽，聪明但多疑)
    - **核心目标/动机**: (他们最想要什么？)
    - **主要冲突/困境**: (什么在阻碍他们？)

    ### 世界设定
    - **时代与地点**: (故事发生在何时何地？是现代都市，还是架空古代，或是未来太空？)
    - **关键地点描述**: (描述几个故事中会反复出现的关键场景，例如：主角的家、神秘的森林、繁华的都市中心等。)
    - **社会规则/文化背景**: (这个世界有什么独特的法律、习俗、信仰或技术水平？)
    - **氛围与基调**: (故事是光明的、黑暗的、悬疑的，还是幽默的？)

    ### 关键情节节点
    根据大纲，列出几个将推动故事发展的关键转折点或核心事件。

    请确保内容详尽、逻辑自洽，为AI后续创作提供清晰、无歧义的指导。
    """,
    "extract_worldview_updates": """
    请你扮演一位严谨的小说档案管理员。
    你的任务是只从“最新创作的章节内容”中，提取出所有**新增的或发生显著变化**的世界观信息。

    **1. 已有的世界观设定 (仅供参考，不要重复输出已知信息)**:
    ```
    {existing_worldview}
    ```

    **2. 最新创作的章节内容**:
    ```
    {new_content}
    ```

    **你的工作流程**:
    1.  仔细阅读“最新创作的章节内容”。
    2.  识别其中出现的任何**新信息**，例如：
        - **新登场的人物**: 姓名、身份、性格等关键信息。
        - **对现有人物状态的重大更新**: 目标改变、获得新能力、关系发生重大变化等。
        - **新出现的地点、物品或概念**。
        - **新揭示的世界规则或背景故事**。
    3.  将这些新信息或变化，以简洁的Markdown列表形式输出。**只输出新增和变化的部分，不要重复世界观中已有的内容。**

    **输出示例**:
    - **林渊**: 在第九章遇到了影璃，确认了她是命格守护者之一，并得知了“九渊之地”是下一步的目标。
    - **新登场人物 - 影璃**: 命格守护者之一，引导林渊，告诉他“无命”是起点，并指引他前往“九渊之地”。
    - **新地点 - 枯树林**: 命格之塔的残骸所在地。

    现在，请开始提取章节《{section_title}》中的世界观更新点。
    """
}

# --- Web App Setup ---
app = Flask(__name__)
log_queue = queue.Queue()
task_manager = None 
token_tracker = None

def log_message(message, task_id=None):
    """Logs a message to the queue, optionally associating it with a task ID."""
    log_entry = {"timestamp": time.time(), "message": message, "task_id": task_id}
    log_queue.put(log_entry)

# --- Core Logic Classes (V5) ---

class TokenTracker:
    def __init__(self, filepath):
        self.filepath = filepath
        self.lock = threading.Lock()
        self.stats = self._load()

    def _load(self):
        try:
            if os.path.exists(self.filepath):
                with open(self.filepath, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except (IOError, json.JSONDecodeError):
            pass
        return {"total_prompt_tokens": 0, "total_completion_tokens": 0, "daily_stats": {}}

    def _save(self):
        with open(self.filepath, 'w', encoding='utf-8') as f:
            json.dump(self.stats, f, indent=2)

    def update(self, prompt_tokens, completion_tokens):
        with self.lock:
            today = datetime.utcnow().strftime('%Y-%m-%d')
            self.stats['total_prompt_tokens'] += prompt_tokens
            self.stats['total_completion_tokens'] += completion_tokens
            
            if today not in self.stats['daily_stats']:
                self.stats['daily_stats'][today] = {"prompt_tokens": 0, "completion_tokens": 0}
            
            self.stats['daily_stats'][today]['prompt_tokens'] += prompt_tokens
            self.stats['daily_stats'][today]['completion_tokens'] += completion_tokens
            
            self._save()
    
    def get_stats(self):
        with self.lock:
            return self.stats

class ContentGenerator:
    """Handles all interactions with the LLM API."""
    def __init__(self, api_url, api_headers, model_id, tracker):
        self.api_url = api_url
        self.api_headers = api_headers
        self.model_id = model_id
        self.tracker = tracker

    def _call_llm(self, prompt, max_tokens=8192, temperature=0.75, task_id=None):
        log_message(f"  > Calling LLM API...", task_id)
        payload = { "model": self.model_id, "messages": [{"role": "system", "content": "You are a creative and intelligent assistant."}, {"role": "user", "content": prompt}], "max_tokens": max_tokens, "temperature": temperature }
        try:
            response = requests.post(self.api_url, headers=self.api_headers, data=json.dumps(payload), timeout=600)
            response.raise_for_status()
            data = response.json()
            content = data['choices'][0]['message']['content']
            
            # Track token usage
            usage = data.get('usage', {})
            prompt_tokens = usage.get('prompt_tokens', 0)
            completion_tokens = usage.get('completion_tokens', 0)
            if self.tracker and (prompt_tokens > 0 or completion_tokens > 0):
                self.tracker.update(prompt_tokens, completion_tokens)
            
            log_message(f"  > LLM response received. Tokens used: P={prompt_tokens}, C={completion_tokens}", task_id)
            return content
        except requests.exceptions.RequestException as e:
            log_message(f"  [ERROR] API request failed: {e}", task_id)
            return f"API_REQUEST_ERROR: {e}"
        except (KeyError, IndexError) as e:
            log_message(f"  [ERROR] Failed to parse API response: {e}", task_id)
            return f"API_RESPONSE_ERROR: {e}"

    def generate_outline(self, topic, outline_type, structure, task_id=None):
        log_message(f"--- Starting AI Outline Generation for '{topic}' ---", task_id)
        prompt_template = OUTLINE_PROMPT_TEMPLATES[outline_type]['default']
        if outline_type == 'novel':
            prompt = prompt_template.format(topic=topic, chapters=structure.get('chapters', 5))
        else: # course
            prompt = prompt_template.format(topic=topic, chapters=structure.get('chapters', 5), sections=structure.get('sections', 4))
        outline = self._call_llm(prompt, task_id=task_id)
        log_message("--- Outline generated. ---", task_id)
        return outline

    def generate_initial_worldview(self, topic, outline, task_id=None):
        log_message(f"--- Generating Initial Worldview for '{topic}' ---", task_id)
        prompt = CONTENT_PROMPT_TEMPLATES['initial_worldview'].format(topic=topic, outline=outline)
        worldview = self._call_llm(prompt, task_id=task_id)
        log_message("--- Initial Worldview generated. ---", task_id)
        return worldview

    def generate_chapter_script(self, worldview, previous_chapter_summary, section_title, task_id=None):
        log_message(f"  > Generating Mind Library script for '{section_title}'...", task_id)
        prompt = CONTENT_PROMPT_TEMPLATES['generate_chapter_script'].format(
            worldview=worldview,
            previous_chapter_summary=previous_chapter_summary,
            section_title=section_title
        )
        script_str = self._call_llm(prompt, max_tokens=2048, task_id=task_id)
        
        json_match = re.search(r'\{.*\}', script_str, re.DOTALL)
        if not json_match:
            log_message("  [ERROR] Failed to find valid JSON object in chapter script response.", task_id)
            return None
        
        try:
            script_obj = json.loads(json_match.group(0))
            log_message(f"  > Script for '{section_title}' generated successfully.", task_id)
            return script_obj
        except json.JSONDecodeError as e:
            log_message(f"  [ERROR] Failed to parse chapter script JSON: {e}", task_id)
            return None

    def extract_worldview_updates(self, existing_worldview, new_content, section_title, task_id=None):
        log_message("  > Extracting worldview updates from new content...", task_id)
        context_worldview = (existing_worldview if len(existing_worldview) < 4000 else '...' + existing_worldview[-4000:])
        prompt = CONTENT_PROMPT_TEMPLATES['extract_worldview_updates'].format(
            existing_worldview=context_worldview, 
            new_content=new_content,
            section_title=section_title
        )
        updates = self._call_llm(prompt, max_tokens=2048, task_id=task_id)
        log_message("  > Worldview updates extracted successfully.", task_id)
        return updates

    def _generate_novel_section(self, task, chapter_plan_obj):
        task_id = task['id']
        log_message(f"--- Generating novel chapter: '{task['section_title']}' ---", task_id)

        if not chapter_plan_obj:
            log_message(f"  [ERROR] Chapter plan was not provided for '{task['section_title']}'.", task_id)
            return f"# {task['topic']}\n\n## {task['section_title']}\n\n*Content generation failed: Chapter plan not provided.*"
        
        chapter_plan = json.dumps(chapter_plan_obj, ensure_ascii=False, indent=2)

        previous_chapter_context = "这是小说的第一章。"
        if task.get('previous_content'):
            log_message("  > Summarizing previous chapter for context...", task_id)
            summary_prompt = CONTENT_PROMPT_TEMPLATES['summarizer'].format(text_to_summarize=task['previous_content'])
            summary = self._call_llm(summary_prompt, max_tokens=500, task_id=task_id)
            if "API_" not in summary:
                previous_chapter_context = summary

        worldview_context = "无"
        worldview_path = os.path.join(task['base_dir'], WORLDVIEW_DIR_NAME, WORLDVIEW_FILE_NAME)
        if os.path.exists(worldview_path):
            log_message("  > Loading worldview context...", task_id)
            with open(worldview_path, 'r', encoding='utf-8') as f:
                worldview_context = f.read()

        log_message(f"  > Generating content for '{task['section_title']}' based on Mind Library plan...", task_id)
        prompt = CONTENT_PROMPT_TEMPLATES['novel'].format(
            topic=task['topic'],
            section_title=task['section_title'],
            word_count=task.get('word_count', 2500),
            previous_chapter_context=previous_chapter_context,
            worldview_context=worldview_context,
            chapter_plan=chapter_plan
        )
        content = self._call_llm(prompt, task_id=task_id)
        if "API_" in content:
            return f"# {task['topic']}\n\n## {task['section_title']}\n\n*Content generation failed.*\n\nError: {content}"
        
        formatted_content = f"# {task['topic']}\n\n## {task['section_title']}\n\n{content}"

        log_message(f"  > Content generation successful for '{task['section_title']}'!", task_id)
        time.sleep(1)
        return formatted_content

    def generate_content_for_section(self, task, chapter_plan_obj=None):
        if task['outline_type'] == 'novel':
            return self._generate_novel_section(task, chapter_plan_obj)
        
        task_id = task['id']
        log_message(f"  > Generating content for '{task['section_title']}'...", task_id)
        prompt_template = CONTENT_PROMPT_TEMPLATES.get(task['outline_type'])
        prompt = prompt_template.format(topic=task['topic'], section_title=task['section_title'])
        content = self._call_llm(prompt, task_id=task_id)
        if "API_" in content:
            return f"# {task['section_title']}\n\n*Content generation failed.*\n\nError: {content}"
        log_message(f"  > Content generation successful!", task_id)
        time.sleep(1)
        return content


class Task:
    """A class to represent a single generation task."""
    def __init__(self, topic, outline, outline_type, word_count=None):
        self.id = str(uuid.uuid4())
        self.topic = topic
        self.outline = outline
        self.outline_type = outline_type
        self.word_count = word_count
        self.status = "queued"
        self.is_cancellation_requested = False
        self.base_dir = os.path.join(WORKS_LIBRARY_PATH, self.outline_type, self.sanitize_filename(self.topic))
        self.sub_tasks = self._create_sub_tasks()
        if not self.sub_tasks: raise ValueError("Failed to parse outline or no sub-tasks were created.")

    @staticmethod
    def sanitize_filename(name, max_length=100):
        name = re.sub(r'^\d+\.\s*', '', name)
        name = name.replace('**', '')
        sanitized_name = re.sub(r'[\\/*?:"<>|]', '_', name)
        return sanitized_name[:max_length]

    def _create_sub_tasks(self):
        sub_tasks = []
        current_chapter_dir = ""
        
        if self.outline_type == 'novel':
            chapter_counter = 1
            for line in self.outline.strip().split('\n'):
                line = line.strip()
                if not line or line.lower().startswith("书名："):
                    continue
                if line.startswith('###'):
                    chapter_title = line.replace('###', '').strip()
                    file_name = f"第{chapter_counter:03d}章-{self.sanitize_filename(chapter_title)}.md"
                    sub_tasks.append({
                        "id": self.id, "topic": self.topic, "outline_type": self.outline_type, 
                        "section_title": chapter_title, 
                        "file_path": os.path.join(self.base_dir, file_name), 
                        "base_dir": self.base_dir, "word_count": self.word_count
                    })
                    chapter_counter += 1
        else: # Course logic
            for line in self.outline.strip().split('\n'):
                line = line.strip()
                if not line or line.lower().startswith("标题："):
                    continue
                if line.startswith('###'):
                    chapter_title = self.sanitize_filename(line.replace('###', '').strip())
                    current_chapter_dir = os.path.join(self.base_dir, chapter_title)
                elif re.match(r'^\d+\.', line) and current_chapter_dir:
                    section_title = line.strip()
                    file_name = self.sanitize_filename(section_title) + ".md"
                    sub_tasks.append({
                        "id": self.id, "topic": self.topic, "outline_type": self.outline_type, 
                        "section_title": section_title, 
                        "file_path": os.path.join(current_chapter_dir, file_name), 
                        "base_dir": self.base_dir, "word_count": self.word_count
                    })
        return sub_tasks

class TaskManager:
    """Manages the queue and execution of generation tasks."""
    def __init__(self, generator):
        self.task_queue = queue.Queue()
        self.generator = generator
        self.is_running = False
        self.thread = None
        self.current_task = None
        self.lock = threading.Lock()
        self.tasks = {}

    def worker(self):
        while self.is_running:
            try:
                task = self.task_queue.get(timeout=1)
                with self.lock:
                    self.current_task = task
                    self.current_task.status = "running"
                log_message(f"--- Starting task: {task.topic} (ID: {task.id}) ---", task.id)
                
                mind_library = []
                mind_library_path = None
                if task.outline_type == 'novel':
                    mind_library_dir = os.path.join(task.base_dir, MIND_LIBRARY_DIR_NAME)
                    if not os.path.exists(mind_library_dir): os.makedirs(mind_library_dir)
                    mind_library_path = os.path.join(mind_library_dir, MIND_LIBRARY_FILE_NAME)
                    if os.path.exists(mind_library_path):
                        with open(mind_library_path, 'r', encoding='utf-8') as f:
                            mind_library = json.load(f)

                previous_content = None
                previous_summary = "这是小说的第一章。"
                for index, sub_task in enumerate(task.sub_tasks):
                    if task.is_cancellation_requested:
                        task.status = "cancelled"
                        log_message(f"Task '{task.topic}' cancelled.", task.id)
                        break
                    
                    if not os.path.exists(os.path.dirname(sub_task['file_path'])): 
                        os.makedirs(os.path.dirname(sub_task['file_path']))
                    
                    if os.path.exists(sub_task['file_path']):
                        log_message(f"  - [Skipping] File exists: {sub_task['file_path']}", task.id)
                        if task.outline_type == 'novel':
                            with open(sub_task['file_path'], 'r', encoding='utf-8') as f: 
                                previous_content = f.read()
                        continue
                    
                    sub_task['previous_content'] = previous_content
                    
                    chapter_plan_obj = None
                    if task.outline_type == 'novel':
                        # Just-in-time script generation
                        if index >= len(mind_library):
                            worldview_path = os.path.join(task.base_dir, WORLDVIEW_DIR_NAME, WORLDVIEW_FILE_NAME)
                            worldview_context = ""
                            if os.path.exists(worldview_path):
                                with open(worldview_path, 'r', encoding='utf-8') as f:
                                    worldview_context = f.read()
                            
                            if previous_content:
                                summary = self.generator._call_llm(CONTENT_PROMPT_TEMPLATES['summarizer'].format(text_to_summarize=previous_content), max_tokens=500, task_id=task.id)
                                if "API_" not in summary: previous_summary = summary

                            new_script = self.generator.generate_chapter_script(worldview_context, previous_summary, sub_task['section_title'], task.id)
                            if new_script:
                                mind_library.append(new_script)
                                with open(mind_library_path, 'w', encoding='utf-8') as f:
                                    json.dump(mind_library, f, ensure_ascii=False, indent=2)
                                chapter_plan_obj = new_script
                            else:
                                log_message(f"  [ERROR] Failed to generate script for '{sub_task['section_title']}'. Skipping chapter.", task.id)
                                continue
                        else:
                            chapter_plan_obj = mind_library[index]
                    
                    content = self.generator.generate_content_for_section(sub_task, chapter_plan_obj)
                    
                    try:
                        with open(sub_task['file_path'], 'w', encoding='utf-8') as f: 
                            f.write(content)
                        log_message(f"    - [Success] Saved file: {sub_task['file_path']}", task.id)
                        previous_content = content
                        
                        if task.outline_type == 'novel' and "generation failed" not in content:
                            worldview_path = os.path.join(task.base_dir, WORLDVIEW_DIR_NAME, WORLDVIEW_FILE_NAME)
                            if os.path.exists(worldview_path):
                                with open(worldview_path, 'r', encoding='utf-8') as f: 
                                    existing_worldview = f.read()
                                
                                updates = self.generator.extract_worldview_updates(existing_worldview, content, sub_task['section_title'], task.id)
                                
                                if "API_" not in updates and updates.strip():
                                    with open(worldview_path, 'a', encoding='utf-8') as f:
                                        f.write(f"\n\n---\n\n### 《{sub_task['section_title']}》章节更新\n\n{updates}")
                                    log_message(f"    - [Success] Appended updates to Worldview.", task.id)
                            
                            if mind_library and mind_library_path and index < len(mind_library):
                                mind_library[index]["status"] = "completed"
                                with open(mind_library_path, 'w', encoding='utf-8') as f:
                                    json.dump(mind_library, f, ensure_ascii=False, indent=2)
                                log_message(f"    - [Success] Updated Mind Library status for '{sub_task['section_title']}'", task.id)

                    except IOError as e:
                        log_message(f"    - [Failed] Error writing file: {e}", task.id)
                
                if task.status not in ["cancelled", "failed"]: 
                    task.status = "completed"
                log_message(f"--- Task '{task.topic}' {task.status} ---", task.id)
                
                with self.lock:
                    self.tasks[task.id] = task
                    self.current_task = None
                self.task_queue.task_done()
            except queue.Empty: 
                continue
        log_message("Worker thread has stopped.")

    def start(self):
        if not self.is_running:
            self.is_running = True
            self.thread = threading.Thread(target=self.worker)
            self.thread.daemon = True
            self.thread.start()
            log_message("Task manager started.")

    def add_task(self, task):
        with self.lock: self.tasks[task.id] = task
        self.task_queue.put(task)
        log_message(f"Task '{task.topic}' added to queue.", task.id)

    def cancel_task(self, task_id):
        with self.lock:
            task = self.tasks.get(task_id)
            if not task: return False, "Task not found."
            if task.status == "running":
                task.is_cancellation_requested = True
                return True, "Cancellation requested."
            if task.status == "queued":
                task.status = "cancelled"
                new_q = queue.Queue()
                while not self.task_queue.empty():
                    item = self.task_queue.get()
                    if item.id != task_id: new_q.put(item)
                self.task_queue = new_q
                log_message(f"Task '{task.topic}' removed from queue.", task.id)
                return True, "Task removed from queue."
            return False, f"Task is already {task.status}."

    def get_tasks_status(self):
        with self.lock:
            status_list = []
            if self.current_task: status_list.append({"id": self.current_task.id, "topic": self.current_task.topic, "status": self.current_task.status})
            for task in list(self.task_queue.queue): status_list.append({ "id": task.id, "topic": task.topic, "status": task.status })
            for task in self.tasks.values():
                if task.status in ['completed', 'cancelled', 'failed'] and task.id not in [t['id'] for t in status_list]: status_list.append({ "id": task.id, "topic": task.topic, "status": task.status })
            return status_list

# --- HTML Template (V5 - Added Token Stats) ---
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI Content Factory v7.0</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css" rel="stylesheet">
    <style>
        :root { --main-bg: #f7fafc; --sidebar-bg: #ffffff; --text-color: #2d3748; --accent-color: #4299e1; --font-serif: 'KaiTi', 'STKaiti', serif; }
        body { font-family: 'Inter', sans-serif; background-color: var(--main-bg); color: var(--text-color); }
        .sidebar-icon { width: 20px; text-align: center; }
        .loader { border-top-color: var(--accent-color); }
        .nav-item.active { background-color: #ebf8ff; color: #2b6cb0; border-right: 4px solid var(--accent-color); font-weight: 600; }
        .prose-styles { max-width: 100%; font-family: var(--font-serif); line-height: 2; font-size: 1.15em; }
        .prose-styles h1, .prose-styles h2, .prose-styles h3 { font-weight: 700; margin-top: 1.5em; margin-bottom: 0.8em; font-family: 'Inter', sans-serif; }
        .showcase-bg, .preview-bg {
            background-color: #fdfcf8;
            border: 1px solid #e2d9c8;
            box-shadow: 0 4px 12px rgba(0,0,0,0.08);
            color: #3a352d;
        }
        #worldview-content-cards .card, #mind-library-list .card, .stat-card {
            background: white;
            border-radius: 8px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            transition: transform 0.2s, box-shadow 0.2s;
        }
        #worldview-content-cards .card:hover, #mind-library-list .card:hover { transform: translateY(-5px); box-shadow: 0 10px 15px rgba(0,0,0,0.1); }
        .character-avatar {
            width: 80px;
            height: 80px;
            border-radius: 50%;
            object-fit: cover;
            border: 3px solid #e2e8f0;
        }
    </style>
</head>
<body class="flex h-screen">

    <!-- Sidebar -->
    <aside class="w-64 bg-white shadow-md flex flex-col flex-shrink-0">
        <div class="p-4 border-b flex justify-between items-center">
            <h2 class="text-2xl font-bold text-center">
                <i class="fas fa-brain text-blue-500"></i>
                <span data-translate-key="app_title">AI Factory</span>
            </h2>
            <select id="lang-switcher" class="bg-gray-100 border-gray-300 rounded-md text-sm p-1">
                <option value="en">EN</option>
                <option value="zh">中文</option>
            </select>
        </div>
        <nav class="flex-grow">
            <a href="#" id="nav-creator" class="nav-item flex items-center space-x-3 p-4 text-lg font-medium hover:bg-gray-100 active"><i class="fas fa-plus-circle sidebar-icon"></i><span data-translate-key="nav_creator">Create Project</span></a>
            <a href="#" id="nav-library" class="nav-item flex items-center space-x-3 p-4 text-lg font-medium hover:bg-gray-100"><i class="fas fa-book-open sidebar-icon"></i><span data-translate-key="nav_library">Project Library</span></a>
            <a href="#" id="nav-worldview" class="nav-item flex items-center space-x-3 p-4 text-lg font-medium hover:bg-gray-100"><i class="fas fa-globe-americas sidebar-icon"></i><span data-translate-key="nav_worldview">Worldview</span></a>
            <a href="#" id="nav-mind-library" class="nav-item flex items-center space-x-3 p-4 text-lg font-medium hover:bg-gray-100"><i class="fas fa-clipboard-list sidebar-icon"></i><span data-translate-key="nav_mind_library">Mind Library</span></a>
            <a href="#" id="nav-queue" class="nav-item flex items-center space-x-3 p-4 text-lg font-medium hover:bg-gray-100"><i class="fas fa-tasks sidebar-icon"></i><span data-translate-key="nav_queue">Task Queue</span></a>
            <a href="#" id="nav-showcase" class="nav-item flex items-center space-x-3 p-4 text-lg font-medium hover:bg-gray-100"><i class="fas fa-magic sidebar-icon"></i><span data-translate-key="nav_showcase">Showcase</span></a>
            <a href="#" id="nav-stats" class="nav-item flex items-center space-x-3 p-4 text-lg font-medium hover:bg-gray-100"><i class="fas fa-chart-line sidebar-icon"></i><span data-translate-key="nav_stats">Token Stats</span></a>
        </nav>
        <div class="p-4 border-t">
            <h3 class="text-lg font-semibold mb-2" data-translate-key="logs_title">Logs</h3>
            <div id="log-panel" class="h-48 bg-gray-900 text-white font-mono text-xs rounded-lg p-2 overflow-y-auto"></div>
        </div>
    </aside>

    <!-- Main Content -->
    <main class="flex-1 p-8 overflow-y-auto bg-gray-50">
        <!-- All Views Here... -->
        <div id="view-creator" class="view-panel">
            <h1 class="text-3xl font-bold mb-6" data-translate-key="creator_title">Create New Project</h1>
            <form id="main-form" class="space-y-6 bg-white p-8 rounded-lg shadow-lg">
                <div>
                    <label class="block text-lg font-medium text-gray-700 mb-2" data-translate-key="creator_topic_label">1. Topic</label>
                    <input type="text" name="topic" id="topic" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm p-3" data-translate-key-placeholder="creator_topic_placeholder">
                </div>
                <div class="grid grid-cols-1 md:grid-cols-2 gap-6">
                    <div>
                        <label for="outline_type" class="block text-sm font-medium text-gray-700" data-translate-key="creator_type_label">Project Type</label>
                        <select id="outline_type" name="outline_type" class="mt-1 block w-full rounded-md border-gray-300 p-2">
                            <option value="novel" data-translate-key="option_novel">Novel</option>
                            <option value="course" data-translate-key="option_course">Course</option>
                        </select>
                    </div>
                    <div id="novel-options">
                        <label for="word_count" class="block text-sm font-medium text-gray-700" data-translate-key="creator_word_count_label">Chapter Word Count</label>
                        <input type="number" name="word_count" id="word_count" value="2500" class="mt-1 block w-full rounded-md border-gray-300 p-2">
                    </div>
                </div>
                <div class="grid grid-cols-1 md:grid-cols-2 gap-6">
                    <div>
                        <label for="chapters" class="block text-sm font-medium text-gray-700" data-translate-key="creator_chapters_label">Number of Chapters</label>
                        <input type="number" name="chapters" id="chapters" value="10" class="mt-1 block w-full rounded-md border-gray-300 p-2">
                    </div>
                    <div id="sections-container">
                        <label for="sections" class="block text-sm font-medium text-gray-700" data-translate-key="creator_sections_label">Sections per Chapter</label>
                        <input type="number" name="sections" id="sections" value="4" class="mt-1 block w-full rounded-md border-gray-300 p-2">
                    </div>
                </div>
                <button type="button" id="generate-outline-btn" class="w-full flex justify-center py-3 px-4 rounded-md shadow-sm text-white bg-blue-600 hover:bg-blue-700">
                    <i class="fas fa-lightbulb mr-2"></i><span data-translate-key="creator_generate_btn">Generate Outline</span>
                </button>
            </form>
            <div id="outline-preview-container" class="mt-8 hidden bg-white p-8 rounded-lg shadow-lg">
                <h3 class="text-xl font-semibold mb-4" data-translate-key="creator_preview_title">Outline Preview</h3>
                <div id="outline-preview" class="p-4 bg-gray-50 border rounded-md max-h-96 overflow-y-auto whitespace-pre-wrap font-mono"></div>
                <button type="button" id="start-generation-btn" class="mt-6 w-full flex justify-center py-3 px-4 rounded-md shadow-sm text-white bg-green-600 hover:bg-green-700">
                    <i class="fas fa-play-circle mr-2"></i><span data-translate-key="creator_confirm_btn">Confirm and Start Generation</span>
                </button>
            </div>
        </div>
        <div id="view-library" class="view-panel hidden">
            <h1 class="text-3xl font-bold mb-6" data-translate-key="library_title">Project Library</h1>
            <div class="grid grid-cols-1 md:grid-cols-2 lg:grid-cols-3 gap-6" id="library-grid"></div>
        </div>
        <div id="view-worldview" class="view-panel hidden">
            <h1 class="text-3xl font-bold mb-6" data-translate-key="worldview_title">Novel Worldview</h1>
            <div class="bg-white p-8 rounded-lg shadow-lg">
                <div class="mb-6">
                    <label for="worldview-project-select" class="block text-sm font-medium text-gray-700" data-translate-key="worldview_select_novel">Select Novel</label>
                    <select id="worldview-project-select" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm p-2"></select>
                </div>
                <div id="worldview-display-container" class="hidden space-y-8">
                    <div id="worldview-content-cards"></div>
                </div>
            </div>
        </div>
        <div id="view-mind-library" class="view-panel hidden">
            <h1 class="text-3xl font-bold mb-6" data-translate-key="mind_library_title">Novel Mind Library</h1>
            <div class="bg-white p-8 rounded-lg shadow-lg">
                <div class="mb-6">
                    <label for="mind-library-project-select" class="block text-sm font-medium text-gray-700" data-translate-key="mind_library_select_novel">Select Novel</label>
                    <select id="mind-library-project-select" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm p-2"></select>
                </div>
                <div id="mind-library-display-container" class="hidden">
                    <div id="mind-library-list" class="space-y-4 max-h-[70vh] overflow-y-auto pr-4"></div>
                </div>
            </div>
        </div>
        <div id="view-preview" class="view-panel hidden">
            <button id="back-to-library-btn" class="mb-6 bg-gray-200 hover:bg-gray-300 text-gray-800 font-bold py-2 px-4 rounded inline-flex items-center">
                <i class="fas fa-arrow-left mr-2"></i><span data-translate-key="preview_back_btn">Back to Library</span>
            </button>
            <div class="flex flex-col lg:flex-row space-y-6 lg:space-y-0 lg:space-x-6">
                <div class="w-full lg:w-1/3 bg-white p-4 rounded-lg shadow-lg flex-shrink-0">
                    <h2 id="preview-title" class="text-xl font-bold mb-4 border-b pb-2 flex justify-between items-center"></h2>
                    <div id="preview-file-tree" class="max-h-[70vh] overflow-y-auto"></div>
                </div>
                <div class="w-full lg:w-2/3 bg-white p-6 rounded-lg shadow-lg preview-bg">
                    <div id="preview-content" class="prose-styles max-h-[75vh] overflow-y-auto p-4"></div>
                </div>
            </div>
        </div>
        <div id="view-queue" class="view-panel hidden">
            <h1 class="text-3xl font-bold mb-6" data-translate-key="queue_title">Task Queue</h1>
            <div id="task-list" class="bg-white p-6 rounded-lg shadow-lg space-y-4"></div>
        </div>
        <div id="view-showcase" class="view-panel hidden">
            <h1 class="text-3xl font-bold mb-6" data-translate-key="showcase_title">AI Writing Showcase</h1>
            <div class="bg-white p-8 rounded-lg shadow-lg">
                <div class="grid grid-cols-1 md:grid-cols-3 gap-4 mb-6 items-end">
                    <div class="md:col-span-1">
                        <label for="showcase-project-select" class="block text-sm font-medium text-gray-700" data-translate-key="showcase_select_novel">Select Novel</label>
                        <select id="showcase-project-select" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm p-2"></select>
                    </div>
                     <div class="md:col-span-1">
                        <label for="showcase-section-select" class="block text-sm font-medium text-gray-700" data-translate-key="showcase_select_section">Section</label>
                        <select id="showcase-section-select" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm p-2" disabled></select>
                    </div>
                    <div class="md:col-span-1 flex space-x-2">
                        <button id="start-showcase-btn" class="w-full flex items-center justify-center py-2 px-4 rounded-md shadow-sm text-white bg-purple-600 hover:bg-purple-700 disabled:bg-gray-400" disabled>
                            <i class="fas fa-play mr-2"></i><span data-translate-key="showcase_start_btn">Start</span>
                        </button>
                        <button id="stop-showcase-btn" class="w-full flex items-center justify-center py-2 px-4 rounded-md shadow-sm text-white bg-red-600 hover:bg-red-700 hidden">
                            <i class="fas fa-stop mr-2"></i><span data-translate-key="showcase_stop_btn">Stop</span>
                        </button>
                    </div>
                </div>
                <div id="showcase-output-wrapper" class="p-8 rounded-lg showcase-bg h-[60vh] overflow-y-auto">
                    <div id="showcase-output" class="prose-styles"></div>
                </div>
            </div>
        </div>
        <div id="view-stats" class="view-panel hidden">
            <h1 class="text-3xl font-bold mb-6" data-translate-key="stats_title">Token Usage Statistics</h1>
            <div class="grid grid-cols-1 md:grid-cols-3 gap-6 mb-6">
                <div class="stat-card p-6 text-center">
                    <h3 class="text-lg font-semibold text-gray-500" data-translate-key="stats_total_prompt">Total Prompt Tokens</h3>
                    <p id="total-prompt-tokens" class="text-4xl font-bold"></p>
                </div>
                <div class="stat-card p-6 text-center">
                    <h3 class="text-lg font-semibold text-gray-500" data-translate-key="stats_total_completion">Total Completion Tokens</h3>
                    <p id="total-completion-tokens" class="text-4xl font-bold"></p>
                </div>
                <div class="stat-card p-6 text-center">
                    <h3 class="text-lg font-semibold text-gray-500" data-translate-key="stats_total">Total Tokens</h3>
                    <p id="total-tokens" class="text-4xl font-bold"></p>
                </div>
            </div>
            <div class="bg-white p-8 rounded-lg shadow-lg">
                <canvas id="token-chart"></canvas>
            </div>
        </div>
        
        <div id="loader" class="hidden fixed inset-0 bg-black bg-opacity-50 flex justify-center items-center z-50">
            <div class="loader ease-linear rounded-full border-8 border-t-8 border-gray-200 h-32 w-32"></div>
            <p id="loader-text" class="text-white text-xl ml-6"></p>
        </div>
    </main>

    <script>
        // --- Translations, State, DOM Elements... (Mostly unchanged)
        const translations = {
            en: {
                app_title: "AI Factory", nav_creator: "Create Project", nav_library: "Project Library", nav_worldview: "Worldview", nav_mind_library: "Mind Library", nav_queue: "Task Queue", nav_showcase: "Showcase", nav_stats: "Token Stats", logs_title: "Logs", creator_title: "Create New Project", creator_topic_label: "1. Topic", creator_topic_placeholder: "e.g., A story about a time-traveling detective", creator_type_label: "Project Type", option_course: "Course", option_novel: "Novel", creator_word_count_label: "Chapter Word Count", creator_chapters_label: "Number of Chapters", creator_sections_label: "Sections per Chapter", creator_generate_btn: "Generate Outline", creator_preview_title: "Outline Preview", creator_confirm_btn: "Confirm and Start Generation", library_title: "Project Library", library_empty: "No projects found. Create one to get started!", worldview_title: "Novel Worldview", worldview_select_novel: "Select Novel", mind_library_title: "Novel Mind Library", mind_library_select_novel: "Select Novel", preview_back_btn: "Back to Library", preview_select_file: "Select a file to view its content.", queue_title: "Task Queue", queue_empty: "The task queue is empty.", showcase_title: "AI Writing Showcase", showcase_select_novel: "Select Novel", showcase_select_section: "Start Chapter", showcase_start_btn: "Start", showcase_stop_btn: "Stop", showcase_loading: "Loading chapters...", showcase_ready: "Select a novel and a starting chapter to begin.", showcase_finished: "Showcase finished.", stats_title: "Token Usage Statistics", stats_total_prompt: "Total Prompt Tokens", stats_total_completion: "Total Completion Tokens", stats_total: "Total Tokens", status_running: "Running", status_queued: "Queued", status_completed: "Completed", status_cancelled: "Cancelled", status_failed: "Failed", action_cancel: "Cancel", action_download_zip: "ZIP", alert_enter_topic: "Please enter a topic!", alert_task_added: "Task successfully added to the queue!", alert_task_creation_failed: "Task creation failed. The outline might be invalid or empty. Please try generating the outline again.", alert_cancel_confirm: "Are you sure you want to cancel this task?", alert_cancel_success: "Task cancellation successful.", loader_generating_outline: "Generating Outline...", loader_generating_worldview: "Generating Worldview...", loader_generating_mind_library: "Generating Mind Library...", loader_adding_task: "Adding task to queue...", loader_loading_library: "Loading project library...", loader_loading_project: "Loading project...", loader_cancelling_task: "Cancelling task...",
            },
            zh: {
                app_title: "AI 作家", nav_creator: "创建项目", nav_library: "作品库", nav_worldview: "世界观", nav_mind_library: "思维库", nav_queue: "任务队列", nav_showcase: "AI 作家", nav_stats: "Token 统计", logs_title: "日志", creator_title: "创建新项目", creator_topic_label: "1. 主题", creator_topic_placeholder: "例如：一个关于时间旅行侦探的故事", creator_type_label: "项目类型", option_course: "课程", option_novel: "小说", creator_word_count_label: "每章字数", creator_chapters_label: "章节数", creator_sections_label: "每章小节数", creator_generate_btn: "生成大纲", creator_preview_title: "大纲预览", creator_confirm_btn: "确认并开始生成", library_title: "作品库", library_empty: "未找到任何项目。请先创建一个。", worldview_title: "小说世界观", worldview_select_novel: "选择小说", mind_library_title: "小说思维库", mind_library_select_novel: "选择小说", preview_back_btn: "返回作品库", preview_select_file: "请选择一个文件以查看其内容。", queue_title: "任务队列", queue_empty: "任务队列是空的。", showcase_title: "AI 写作秀场", showcase_select_novel: "选择小说", showcase_select_section: "起始章节", showcase_start_btn: "开始", showcase_stop_btn: "停止", showcase_loading: "正在加载章节...", showcase_ready: "请选择一部小说和起始章节。", showcase_finished: "演示已结束。", stats_title: "Token 使用统计", stats_total_prompt: "总输入 Tokens", stats_total_completion: "总输出 Tokens", stats_total: "总消耗 Tokens", status_running: "运行中", status_queued: "排队中", status_completed: "已完成", status_cancelled: "已取消", status_failed: "失败", action_cancel: "取消", action_download_zip: "下载ZIP", alert_enter_topic: "请输入核心主题！", alert_task_added: "任务已成功添加到队列！", alert_task_creation_failed: "任务创建失败。生成的大纲可能无效或为空，请尝试重新生成大纲。", alert_cancel_confirm: "您确定要取消这个任务吗？", alert_cancel_success: "任务已成功取消。", loader_generating_outline: "正在生成大纲...", loader_generating_worldview: "正在创建世界观...", loader_generating_mind_library: "正在创建思维库...", loader_adding_task: "正在添加任务到队列...", loader_loading_library: "正在加载作品库...", loader_loading_project: "正在加载项目...", loader_cancelling_task: "正在取消任务...",
            }
        };
        let currentLang = 'en';
        let typewriterTimeout;
        let isShowcaseRunning = false;
        let completedTaskCount = 0;
        let tokenChart = null;

        const views = { creator: document.getElementById('view-creator'), library: document.getElementById('view-library'), worldview: document.getElementById('view-worldview'), mind_library: document.getElementById('view-mind-library'), preview: document.getElementById('view-preview'), queue: document.getElementById('view-queue'), showcase: document.getElementById('view-showcase'), stats: document.getElementById('view-stats') };
        const navLinks = { creator: document.getElementById('nav-creator'), library: document.getElementById('nav-library'), worldview: document.getElementById('nav-worldview'), mind_library: document.getElementById('nav-mind-library'), queue: document.getElementById('nav-queue'), showcase: document.getElementById('nav-showcase'), stats: document.getElementById('nav-stats') };
        const loader = document.getElementById('loader'), loaderText = document.getElementById('loader-text');
        const logPanel = document.getElementById('log-panel');
        const langSwitcher = document.getElementById('lang-switcher');
        
        const worldviewProjectSelect = document.getElementById('worldview-project-select');
        const worldviewDisplayContainer = document.getElementById('worldview-display-container');
        const worldviewContentCards = document.getElementById('worldview-content-cards');
        
        const mindLibraryProjectSelect = document.getElementById('mind-library-project-select');
        const mindLibraryDisplayContainer = document.getElementById('mind-library-display-container');
        const mindLibraryList = document.getElementById('mind-library-list');

        const showcaseProjectSelect = document.getElementById('showcase-project-select');
        const showcaseSectionSelect = document.getElementById('showcase-section-select');
        const startShowcaseBtn = document.getElementById('start-showcase-btn');
        const stopShowcaseBtn = document.getElementById('stop-showcase-btn');
        const showcaseOutput = document.getElementById('showcase-output');

        function setLanguage(lang) {
            currentLang = lang;
            localStorage.setItem('ai-factory-lang', lang);
            langSwitcher.value = lang;
            document.documentElement.lang = lang;
            document.querySelectorAll('[data-translate-key]').forEach(el => {
                const key = el.getAttribute('data-translate-key');
                el.innerText = translations[lang][key] || el.innerText;
            });
            document.querySelectorAll('[data-translate-key-placeholder]').forEach(el => {
                const key = el.getAttribute('data-translate-key-placeholder');
                el.placeholder = translations[lang][key] || el.placeholder;
            });
        }

        langSwitcher.addEventListener('change', (e) => setLanguage(e.target.value));

        function switchView(viewName) {
            Object.values(views).forEach(v => v.classList.add('hidden'));
            Object.values(navLinks).forEach(l => l.classList.remove('active'));
            views[viewName].classList.remove('hidden');
            navLinks[viewName].classList.add('active');
            stopShowcase();
            if (viewName === 'library') loadLibrary();
            if (viewName === 'worldview') loadWorldviewProjects();
            if (viewName === 'mind_library') loadMindLibraryProjects();
            if (viewName === 'queue') updateTaskQueueView();
            if (viewName === 'showcase') loadShowcaseProjects();
            if (viewName === 'stats') loadTokenStats();
        }

        Object.keys(navLinks).forEach(key => navLinks[key].addEventListener('click', (e) => { e.preventDefault(); switchView(key); }));

        async function apiCall(endpoint, options = {}) {
            try {
                const response = await fetch(endpoint, options);
                if (!response.ok) {
                    const error = await response.json();
                    throw new Error(error.error || `API call failed: ${response.statusText}`);
                }
                return options.raw ? response : response.json();
            } catch (error) {
                console.error(`Error calling ${endpoint}:`, error);
                alert(`An error occurred: ${error.message}`);
                hideLoader();
                return null;
            }
        }
        function showLoader(textKey) { loaderText.innerText = translations[currentLang][textKey]; loader.classList.remove('hidden'); }
        function hideLoader() { loader.classList.add('hidden'); }
        
        // --- Creator View Logic ---
        function toggleCreatorOptions() {
            const type = document.getElementById('outline_type').value;
            document.getElementById('novel-options').style.display = type === 'novel' ? 'block' : 'none';
            document.getElementById('sections-container').style.display = type === 'course' ? 'block' : 'none';
        }
        document.getElementById('outline_type').addEventListener('change', toggleCreatorOptions);
        
        document.getElementById('generate-outline-btn').addEventListener('click', async () => {
            const topic = document.getElementById('topic').value;
            if (!topic) { alert(translations[currentLang].alert_enter_topic); return; }
            showLoader('loader_generating_outline');
            const data = Object.fromEntries(new FormData(document.getElementById('main-form')).entries());
            const result = await apiCall('/generate-outline', { method: 'POST', headers: { 'Content-Type': 'application/json' }, body: JSON.stringify(data) });
            hideLoader();
            if (result) {
                document.getElementById('outline-preview').innerText = result.outline;
                document.getElementById('outline-preview-container').classList.remove('hidden');
            }
        });
        document.getElementById('start-generation-btn').addEventListener('click', async () => {
            const taskData = { 
                original_topic: document.getElementById('topic').value,
                outline: document.getElementById('outline-preview').innerText,
                outline_type: document.getElementById('outline_type').value,
                word_count: document.getElementById('word_count').value
            };
            showLoader('loader_adding_task');
            const result = await apiCall('/start-task', { method: 'POST', headers: { 'Content-Type': 'application/json' }, body: JSON.stringify(taskData) });
            hideLoader();
            if (result) {
                if(result.status === 'success') {
                    alert(translations[currentLang].alert_task_added);
                    document.getElementById('main-form').reset();
                    toggleCreatorOptions();
                    document.getElementById('outline-preview-container').classList.add('hidden');
                    switchView('queue');
                } else {
                    alert(translations[currentLang].alert_task_creation_failed + ` (${result.error})`);
                }
            }
        });

        // --- Library, Worldview, Mind Library, etc. ---
        // (These sections are largely the same as V4, so they are omitted for brevity but are included in the full code)
        async function loadLibrary() {
            showLoader('loader_loading_library');
            const projects = await apiCall('/library');
            hideLoader();
            const libraryGrid = document.getElementById('library-grid');
            libraryGrid.innerHTML = '';
            if (projects && (projects.course.length > 0 || projects.novel.length > 0)) {
                const renderProject = (p, type) => {
                    const icon = type === 'course' ? 'fa-chalkboard-teacher' : 'fa-book';
                    const typeText = type === 'course' ? translations[currentLang].option_course : translations[currentLang].option_novel;
                    return `<div class="bg-white p-6 rounded-lg shadow-md hover:shadow-xl transition-shadow cursor-pointer" onclick="loadProjectPreview('${type}', '${p}')">
                        <div class="flex items-center space-x-4"> <i class="fas ${icon} text-3xl text-blue-500"></i> <div> <h3 class="text-lg font-bold">${p}</h3> <p class="text-sm text-gray-500">${typeText}</p> </div> </div>
                    </div>`;
                };
                projects.course.forEach(p => libraryGrid.innerHTML += renderProject(p, 'course'));
                projects.novel.forEach(p => libraryGrid.innerHTML += renderProject(p, 'novel'));
            } else {
                libraryGrid.innerHTML = `<p class="text-gray-500 col-span-full text-center">${translations[currentLang].library_empty}</p>`;
            }
        }
        async function loadProjectPreview(type, project) {
            showLoader('loader_loading_project');
            const data = await apiCall(`/library/${type}/${project}`);
            hideLoader();
            if (data) {
                document.getElementById('preview-title').innerHTML = `<i class="fas ${type === 'course' ? 'fa-chalkboard-teacher' : 'fa-book'} mr-2"></i> ${data.name} 
                    <a href="/download/zip/${type}/${project}" class="text-sm bg-green-500 hover:bg-green-600 text-white font-bold py-1 px-3 rounded-full ml-4">
                        <i class="fas fa-download mr-1"></i>${translations[currentLang].action_download_zip}
                    </a>`;
                
                let treeHtml = '';
                if (type === 'novel') {
                    treeHtml += `<ul class="pl-4 ml-2">`;
                    data.tree.forEach((file) => {
                         const filepath = encodeURIComponent(file);
                         treeHtml += `<li class="py-1"><a href="#" class="text-gray-600 hover:text-blue-600" onclick="loadFileContent(event, '${type}', '${project}', '${filepath}')">${file.replace('.md','')}</a></li>`;
                    });
                    treeHtml += `</ul>`;
                } else { // Course
                    for (const chapter in data.tree) {
                        treeHtml += `<div class="mt-2"> <h4 class="font-semibold text-gray-700"><i class="fas fa-folder-open mr-2 text-yellow-500"></i>${chapter}</h4> <ul class="pl-4 border-l-2 border-gray-200 ml-2">`;
                        data.tree[chapter].forEach((file) => {
                            const filepath = encodeURIComponent(`${chapter}/${file}`);
                            treeHtml += `<li class="py-1"><a href="#" class="text-gray-600 hover:text-blue-600" onclick="loadFileContent(event, '${type}', '${project}', '${filepath}')">${file.replace('.md','')}</a></li>`;
                        });
                        treeHtml += `</ul></div>`;
                    }
                }

                document.getElementById('preview-file-tree').innerHTML = treeHtml;
                document.getElementById('preview-content').innerHTML = `<span>${translations[currentLang].preview_select_file}</span>`;
                switchView('preview');
            }
        }
        async function loadFileContent(event, type, project, filepath) {
            event.preventDefault();
            document.getElementById('preview-content').innerHTML = 'Loading...';
            const data = await apiCall(`/library/content/${type}/${project}/${filepath}`);
            if (data) { document.getElementById('preview-content').innerHTML = marked.parse(data.content); }
        }
        document.getElementById('back-to-library-btn').addEventListener('click', () => switchView('library'));
        
        function parseWorldview(mdContent) {
            const sections = mdContent.split('### ').slice(1);
            const worldview = {};
            sections.forEach(section => {
                const lines = section.split('\\n');
                const title = lines[0].trim();
                const content = lines.slice(1).join('\\n').trim();
                if (title.includes('主要人物')) {
                    worldview.characters = content.split('---').map(charBlock => {
                        const charData = {};
                        charBlock.trim().split('\\n').forEach(line => {
                            const parts = line.split(':');
                            if (parts.length > 1) {
                                const key = parts[0].replace(/\\*\\*/g, '').trim();
                                const value = parts.slice(1).join(':').trim();
                                charData[key] = value;
                            }
                        });
                        return charData;
                    }).filter(c => c['姓名']);
                } else {
                    worldview[title] = content;
                }
            });
            return worldview;
        }

        function renderWorldview(worldview) {
            let html = '<div class="space-y-6">';
            if (worldview.characters) {
                html += '<div><h3 class="text-2xl font-bold mb-4">主要人物</h3>';
                html += '<div class="grid grid-cols-1 md:grid-cols-2 lg:grid-cols-3 gap-6">';
                worldview.characters.forEach(char => {
                    html += `<div class="card p-6 flex flex-col items-center text-center"><img src="https://i.pravatar.cc/150?u=${encodeURIComponent(char['姓名'])}" alt="Avatar" class="character-avatar mb-4"><h4 class="text-lg font-bold">${char['姓名'] || 'N/A'}</h4><p class="text-sm text-gray-500 mb-2">${char['身份/背景'] || ''}</p></div>`;
                });
                html += '</div></div>';
            }
            for (const [title, content] of Object.entries(worldview)) {
                if (title !== 'characters') {
                    html += `<div class="card p-6"><h3 class="text-xl font-bold mb-2">${title}</h3><div class="prose max-w-none text-sm">${marked.parse(content.replace(/\\n/g, '\\n'))}</div></div>`;
                }
            }
            html += '</div>';
            worldviewContentCards.innerHTML = html;
        }

        async function loadWorldviewProjects() {
            worldviewProjectSelect.innerHTML = `<option value="">---</option>`;
            worldviewDisplayContainer.classList.add('hidden');
            const projects = await apiCall('/library');
            if (projects && projects.novel.length > 0) {
                projects.novel.forEach(p => { worldviewProjectSelect.innerHTML += `<option value="${p}">${p}</option>`; });
            }
        }
        worldviewProjectSelect.addEventListener('change', async (e) => {
            const projectName = e.target.value;
            if (!projectName) { worldviewDisplayContainer.classList.add('hidden'); return; }
            const data = await apiCall(`/api/worldview/${projectName}`);
            if (data && data.content) {
                const parsed = parseWorldview(data.content.replace(/\\n/g, '\\n'));
                renderWorldview(parsed);
                worldviewDisplayContainer.classList.remove('hidden');
            } else {
                 worldviewContentCards.innerHTML = 'No worldview found.';
                 worldviewDisplayContainer.classList.remove('hidden');
            }
        });

        async function loadMindLibraryProjects() {
            mindLibraryProjectSelect.innerHTML = `<option value="">---</option>`;
            mindLibraryDisplayContainer.classList.add('hidden');
            const projects = await apiCall('/library');
            if (projects && projects.novel.length > 0) {
                projects.novel.forEach(p => { mindLibraryProjectSelect.innerHTML += `<option value="${p}">${p}</option>`; });
            }
        }
        mindLibraryProjectSelect.addEventListener('change', async (e) => {
            const projectName = e.target.value;
            if (!projectName) { mindLibraryDisplayContainer.classList.add('hidden'); return; }
            const data = await apiCall(`/api/mind-library/${projectName}`);
            if (data && data.library) {
                renderMindLibrary(data.library);
                mindLibraryDisplayContainer.classList.remove('hidden');
            } else {
                mindLibraryList.innerHTML = 'No mind library found.';
                mindLibraryDisplayContainer.classList.remove('hidden');
            }
        });

        function renderMindLibrary(library) {
            let html = '';
            library.forEach((item, index) => {
                const statusColor = item.status === 'completed' ? 'bg-green-100 border-green-500' : 'bg-yellow-100 border-yellow-500';
                const statusIcon = item.status === 'completed' ? 'fa-check-circle text-green-500' : 'fa-clock text-yellow-500';
                html += `<div class="card p-4 ${statusColor} border-l-4"><div class="flex justify-between items-center"><h4 class="text-lg font-bold text-gray-800">${index + 1}. ${item.chapter_title}</h4><i class="fas ${statusIcon}"></i></div><p class="text-sm text-gray-600 mt-2">${item.summary}</p></div>`;
            });
            mindLibraryList.innerHTML = html;
        }
        
        // --- Token Stats View Logic (New) ---
        async function loadTokenStats() {
            const data = await apiCall('/api/token-stats');
            if (!data) return;

            document.getElementById('total-prompt-tokens').textContent = data.total_prompt_tokens.toLocaleString();
            document.getElementById('total-completion-tokens').textContent = data.total_completion_tokens.toLocaleString();
            document.getElementById('total-tokens').textContent = (data.total_prompt_tokens + data.total_completion_tokens).toLocaleString();

            const dailyStats = data.daily_stats || {};
            const sortedDates = Object.keys(dailyStats).sort();
            
            const labels = sortedDates;
            const promptData = sortedDates.map(date => dailyStats[date].prompt_tokens);
            const completionData = sortedDates.map(date => dailyStats[date].completion_tokens);

            const ctx = document.getElementById('token-chart').getContext('2d');
            if(tokenChart) {
                tokenChart.destroy();
            }
            tokenChart = new Chart(ctx, {
                type: 'bar',
                data: {
                    labels: labels,
                    datasets: [
                        {
                            label: 'Prompt Tokens',
                            data: promptData,
                            backgroundColor: 'rgba(54, 162, 235, 0.5)',
                            borderColor: 'rgba(54, 162, 235, 1)',
                            borderWidth: 1
                        },
                        {
                            label: 'Completion Tokens',
                            data: completionData,
                            backgroundColor: 'rgba(255, 99, 132, 0.5)',
                            borderColor: 'rgba(255, 99, 132, 1)',
                            borderWidth: 1
                        }
                    ]
                },
                options: {
                    scales: {
                        y: { beginAtZero: true },
                        x: { stacked: true }
                    },
                    responsive: true,
                    plugins: {
                        title: { display: true, text: 'Daily Token Usage' }
                    }
                }
            });
        }
        
        // --- Other functions (Queue, Showcase, etc.) ---
        // (These sections are largely the same as V4, so they are omitted for brevity but are included in the full code)
        async function updateTaskQueueView() {
            const data = await apiCall('/tasks');
            const taskList = document.getElementById('task-list');
            taskList.innerHTML = '';
            if (data && data.length > 0) {
                data.forEach(task => {
                    let statusColor = 'bg-gray-400';
                    if (task.status === 'running') statusColor = 'bg-blue-500';
                    if (task.status === 'completed') statusColor = 'bg-green-500';
                    if (task.status === 'cancelled') statusColor = 'bg-orange-500';
                    if (task.status === 'failed') statusColor = 'bg-red-500';
                    const statusText = translations[currentLang][`status_${task.status}`] || task.status;
                    taskList.innerHTML += `<div class="flex items-center justify-between p-3 bg-gray-50 rounded-lg"><div class="flex items-center"><span class="inline-block w-3 h-3 ${statusColor} rounded-full mr-3"></span><div> <span class="font-medium">${task.topic}</span> <span class="text-xs text-gray-500 ml-2">(${statusText})</span> </div></div>${task.status === 'queued' || task.status === 'running' ? `<button class="bg-red-500 hover:bg-red-600 text-white text-xs font-bold py-1 px-2 rounded" onclick="cancelTask('${task.id}')"><i class="fas fa-trash-alt mr-1"></i>${translations[currentLang].action_cancel}</button>` : ''}</div>`;
                });
            } else {
                taskList.innerHTML = `<p class="text-gray-500 text-center">${translations[currentLang].queue_empty}</p>`;
            }
        }
        async function cancelTask(taskId) {
            if (!confirm(translations[currentLang].alert_cancel_confirm)) return;
            showLoader('loader_cancelling_task');
            const result = await apiCall(`/tasks/cancel/${taskId}`, { method: 'POST' });
            hideLoader();
            if (result) {
                alert(translations[currentLang].alert_cancel_success);
                updateTaskQueueView();
            }
        }
        async function loadShowcaseProjects() {
            showcaseProjectSelect.innerHTML = `<option value="">---</option>`;
            showcaseSectionSelect.innerHTML = `<option value="">---</option>`;
            showcaseSectionSelect.disabled = true;
            startShowcaseBtn.disabled = true;
            showcaseOutput.innerHTML = `<span class="text-gray-400">${translations[currentLang].showcase_ready}</span>`;
            const projects = await apiCall('/library');
            if (projects && projects.novel.length > 0) {
                projects.novel.forEach(p => { showcaseProjectSelect.innerHTML += `<option value="${p}">${p}</option>`; });
            }
        }
        showcaseProjectSelect.addEventListener('change', async (e) => {
            const projectName = e.target.value;
            showcaseSectionSelect.innerHTML = `<option value="">---</option>`;
            showcaseSectionSelect.disabled = true;
            startShowcaseBtn.disabled = true;
            if (!projectName) return;
            const projectDetails = await apiCall(`/library/novel/${projectName}`);
            if (projectDetails && projectDetails.tree) {
                projectDetails.tree.forEach(chapterFile => {
                    showcaseSectionSelect.innerHTML += `<option value="${chapterFile}">${chapterFile.replace('.md','')}</option>`;
                });
                showcaseSectionSelect.disabled = false;
            }
        });
        showcaseSectionSelect.addEventListener('change', (e) => { startShowcaseBtn.disabled = !e.target.value; });
        function typeWriter(element, text, onFinished, index = 0) {
            if (!isShowcaseRunning) { element.innerHTML = marked.parse(text); return; }
            if (index < text.length) {
                element.innerHTML = marked.parse(text.substring(0, index + 1));
                element.parentElement.scrollTop = element.parentElement.scrollHeight;
                const delay = (text[index] === '，' || text[index] === '。') ? 150 : 30;
                typewriterTimeout = setTimeout(() => typeWriter(element, text, onFinished, index + 1), delay);
            } else {
                 if (onFinished) onFinished();
            }
        }
        async function runShowcaseQueue(sections) {
            if (!isShowcaseRunning || sections.length === 0) {
                stopShowcase();
                showcaseOutput.innerHTML += `<br/><br/><p class="text-center font-bold">${translations[currentLang].showcase_finished}</p>`;
                return;
            }
            const nextSection = sections.shift();
            typeWriter(showcaseOutput, nextSection.content, () => {
                setTimeout(() => runShowcaseQueue(sections), 2500);
            });
        }
        function stopShowcase() {
            isShowcaseRunning = false;
            clearTimeout(typewriterTimeout);
            startShowcaseBtn.classList.remove('hidden');
            stopShowcaseBtn.classList.add('hidden');
            showcaseProjectSelect.disabled = false;
            showcaseSectionSelect.disabled = false;
        }
        startShowcaseBtn.addEventListener('click', async () => {
            const projectName = showcaseProjectSelect.value;
            const sectionName = showcaseSectionSelect.value;
            if (!projectName || !sectionName) return;
            isShowcaseRunning = true;
            startShowcaseBtn.classList.add('hidden');
            stopShowcaseBtn.classList.remove('hidden');
            showcaseProjectSelect.disabled = true;
            showcaseSectionSelect.disabled = true;
            showcaseOutput.innerHTML = `<span class="text-gray-400">${translations[currentLang].showcase_loading}</span>`;
            const data = await apiCall(`/api/showcase-sequence/${projectName}/${encodeURIComponent(sectionName)}`);
            if (data && data.sections) { runShowcaseQueue(data.sections); } 
            else { showcaseOutput.innerText = 'Could not load section sequence.'; stopShowcase(); }
        });
        stopShowcaseBtn.addEventListener('click', stopShowcase);

        // --- Logging & Auto-Refresh ---
        function appendLog(log) {
            if (logPanel.innerText.startsWith('Welcome!')) { logPanel.innerHTML = ''; }
            const logLine = document.createElement('div');
            logLine.innerHTML = `<span>${new Date(log.timestamp * 1000).toLocaleTimeString()}:</span> <span class="text-gray-300">${log.message}</span>`;
            logPanel.appendChild(logLine);
            logPanel.scrollTop = logPanel.scrollHeight;
        }

        setInterval(async () => {
            const tasks = await apiCall('/tasks');
            if(views.queue.offsetParent !== null) { updateTaskQueueView(); }
            const newCompletedCount = tasks.filter(t => t.status === 'completed').length;
            if (newCompletedCount > completedTaskCount) {
                console.log("New task completed, refreshing views...");
                if (views.library.offsetParent !== null) loadLibrary();
                if (views.showcase.offsetParent !== null) loadShowcaseProjects();
                if (views.worldview.offsetParent !== null) {
                    const selectedProject = worldviewProjectSelect.value;
                    loadWorldviewProjects().then(() => {
                        if (selectedProject) {
                            worldviewProjectSelect.value = selectedProject;
                            worldviewProjectSelect.dispatchEvent(new Event('change'));
                        }
                    });
                }
                if (views.mind_library.offsetParent !== null) {
                    const selectedProject = mindLibraryProjectSelect.value;
                    loadMindLibraryProjects().then(() => {
                        if (selectedProject) {
                            mindLibraryProjectSelect.value = selectedProject;
                            mindLibraryProjectSelect.dispatchEvent(new Event('change'));
                        }
                    });
                }
            }
            completedTaskCount = newCompletedCount;
            const logs = await apiCall('/status');
            if (logs) logs.forEach(appendLog);
        }, 3000);

        // --- Initialization ---
        const savedLang = localStorage.getItem('ai-factory-lang') || 'zh';
        setLanguage(savedLang);
        toggleCreatorOptions();
        switchView('creator');
        logPanel.innerHTML = 'Welcome! Waiting for tasks...';
    </script>
</body>
</html>
"""

# --- Helper function for file conversion ---
def markdown_to_plain_text(md_content):
    """Converts Markdown content to plain text."""
    content_sans_titles = re.sub(r'^# .*\n', '', md_content)
    content_sans_titles = re.sub(r'^## .*\n', '', content_sans_titles)
    html = markdown(content_sans_titles)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text()

def markdown_to_docx(md_content):
    """Converts Markdown content to a DOCX file in memory."""
    document = Document()
    html = markdown(md_content)
    soup = BeautifulSoup(html, "html.parser")
    for element in soup.contents:
        if element.name == 'h1': document.add_heading(element.text, level=1)
        elif element.name == 'h2': document.add_heading(element.text, level=2)
        elif element.name == 'h3': document.add_heading(element.text, level=3)
        elif element.name == 'p': document.add_paragraph(element.text)
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'): document.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'pre': document.add_paragraph(element.text)
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- Flask Routes (V5 - Upgraded) ---

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/generate-outline', methods=['POST'])
def generate_outline_api():
    data = request.json
    generator = ContentGenerator(TARGET_API_URL, API_HEADERS, SERVED_MODEL_IDENTIFIER, token_tracker)
    structure = {"chapters": int(data.get('chapters', 10)), "sections": int(data.get('sections', 4))}
    outline = generator.generate_outline(data['topic'], data['outline_type'], structure)
    if outline and "API_" not in outline:
        return jsonify({"outline": outline})
    return jsonify({"error": "Failed to generate outline"}), 500

@app.route('/start-task', methods=['POST'])
def start_task_api():
    data = request.json
    outline = data.get('outline')
    original_topic = data.get('original_topic')
    outline_type = data.get('outline_type')
    if not all([outline, original_topic, outline_type]):
        return jsonify({"status": "error", "error": "Missing data"}), 400
    
    first_line = outline.split('\n', 1)[0]
    title_match = re.search(r'^(?:标题|书名)：\s*(.+)', first_line.strip(), re.IGNORECASE)
    topic = title_match.group(1).strip() if title_match else Task.sanitize_filename(original_topic, max_length=50)

    if outline_type == 'novel':
        generator = ContentGenerator(TARGET_API_URL, API_HEADERS, SERVED_MODEL_IDENTIFIER, token_tracker)
        log_message(f"Creating Worldview for {topic}...")
        initial_worldview = generator.generate_initial_worldview(topic, outline)
        if "API_" in initial_worldview:
            return jsonify({"status": "error", "error": "Failed to generate initial worldview"}), 500
        
        project_base_dir = os.path.join(WORKS_LIBRARY_PATH, 'novel', Task.sanitize_filename(topic))
        worldview_dir = os.path.join(project_base_dir, WORLDVIEW_DIR_NAME)
        if not os.path.exists(worldview_dir): os.makedirs(worldview_dir)
        with open(os.path.join(worldview_dir, WORLDVIEW_FILE_NAME), 'w', encoding='utf-8') as f:
            f.write(initial_worldview)
        
        # Mind library is now generated just-in-time, so we don't create it here.
        # We just create the directory.
        mind_library_dir = os.path.join(project_base_dir, MIND_LIBRARY_DIR_NAME)
        if not os.path.exists(mind_library_dir): os.makedirs(mind_library_dir)


    try:
        task = Task(topic, outline, outline_type, data.get('word_count'))
        task_manager.add_task(task)
        return jsonify({"status": "success", "message": "Task added to queue", "task_id": task.id})
    except ValueError as e:
        return jsonify({"status": "error", "error": str(e)}), 400

@app.route('/status')
def status_api():
    logs = []
    while not log_queue.empty():
        logs.append(log_queue.get_nowait())
    return jsonify(logs)

@app.route('/tasks')
def get_tasks_api():
    return jsonify(task_manager.get_tasks_status())

@app.route('/tasks/cancel/<task_id>', methods=['POST'])
def cancel_task_api(task_id):
    success, message = task_manager.cancel_task(task_id)
    if success:
        return jsonify({"message": message})
    return jsonify({"error": message}), 400

@app.route('/library')
def get_library_api():
    library = {"course": [], "novel": []}
    if not os.path.exists(WORKS_LIBRARY_PATH):
        os.makedirs(os.path.join(WORKS_LIBRARY_PATH, "course"))
        os.makedirs(os.path.join(WORKS_LIBRARY_PATH, "novel"))
    for type in ["course", "novel"]:
        type_path = os.path.join(WORKS_LIBRARY_PATH, type)
        if os.path.exists(type_path):
            library[type] = sorted([d for d in os.listdir(type_path) if os.path.isdir(os.path.join(type_path, d))])
    return jsonify(library)

@app.route('/library/<type>/<project_name>')
def get_project_details_api(type, project_name):
    project_path = os.path.join(WORKS_LIBRARY_PATH, type, project_name)
    if not os.path.exists(project_path): return jsonify({"error": "Project not found"}), 404
    
    if type == 'novel':
        files = sorted([f for f in os.listdir(project_path) if f.endswith('.md')])
        return jsonify({"name": project_name, "tree": files})
    else: # Course
        tree = {}
        for chapter in sorted([d for d in os.listdir(project_path) if os.path.isdir(os.path.join(project_path, d)) and d != WORLDVIEW_DIR_NAME]):
            chapter_path = os.path.join(project_path, chapter)
            files = sorted([f for f in os.listdir(chapter_path) if f.endswith('.md')])
            tree[chapter] = files
        return jsonify({"name": project_name, "tree": tree})

@app.route('/library/content/<type>/<project_name>/<path:filepath>')
def get_file_content_api(type, project_name, filepath):
    full_path = os.path.join(WORKS_LIBRARY_PATH, type, project_name, filepath)
    if not os.path.exists(full_path): return jsonify({"error": "File not found"}), 404
    with open(full_path, 'r', encoding='utf-8') as f: content = f.read()
    return jsonify({"content": content})

@app.route('/download/zip/<type>/<project_name>')
def download_project_zip_api(type, project_name):
    project_path = os.path.join(WORKS_LIBRARY_PATH, type, project_name)
    if not os.path.exists(project_path): return jsonify({"error": "Project not found"}), 404
    memory_file = BytesIO()
    with shutil.ZipFile(memory_file, 'w', shutil.ZIP_DEFLATED) as zf:
        for root, _, files in os.walk(project_path):
            if WORLDVIEW_DIR_NAME in root or MIND_LIBRARY_DIR_NAME in root: continue
            for file in files:
                zf.write(os.path.join(root, file), os.path.relpath(os.path.join(root, file), os.path.join(project_path, '..')))
    memory_file.seek(0)
    return send_file(memory_file, download_name=f'{project_name}.zip', as_attachment=True)

@app.route('/api/worldview/<project_name>')
def worldview_api(project_name):
    worldview_file = os.path.join(WORKS_LIBRARY_PATH, 'novel', project_name, WORLDVIEW_DIR_NAME, WORLDVIEW_FILE_NAME)
    if os.path.exists(worldview_file):
        with open(worldview_file, 'r', encoding='utf-8') as f: content = f.read()
        return jsonify({"content": content})
    return jsonify({"content": ""})

@app.route('/api/mind-library/<project_name>')
def mind_library_api(project_name):
    mind_library_file = os.path.join(WORKS_LIBRARY_PATH, 'novel', project_name, MIND_LIBRARY_DIR_NAME, MIND_LIBRARY_FILE_NAME)
    if os.path.exists(mind_library_file):
        with open(mind_library_file, 'r', encoding='utf-8') as f: library = json.load(f)
        return jsonify({"library": library})
    return jsonify({"library": []})

@app.route('/api/token-stats')
def token_stats_api():
    return jsonify(token_tracker.get_stats())

@app.route('/api/showcase-sequence/<project_name>/<path:section_name>')
def get_showcase_sequence_api(project_name, section_name):
    project_path = os.path.join(WORKS_LIBRARY_PATH, 'novel', project_name)
    if not os.path.exists(project_path): return jsonify({"error": "Project not found"}), 404
    
    all_chapters = sorted([f for f in os.listdir(project_path) if f.endswith('.md')])
    try:
        start_chapter_index = all_chapters.index(section_name)
    except ValueError:
        return jsonify({"error": "Starting chapter not found"}), 404
        
    sections_to_play = []
    for chapter_idx in range(start_chapter_index, len(all_chapters)):
        current_chapter_name = all_chapters[chapter_idx]
        with open(os.path.join(project_path, current_chapter_name), 'r', encoding='utf-8') as f: 
            content = f.read()
        sections_to_play.append({"name": current_chapter_name, "content": content})
    return jsonify({"sections": sections_to_play})


# --- Main Execution ---
if __name__ == '__main__':
    for dir_path in [WORKS_LIBRARY_PATH, os.path.join(WORKS_LIBRARY_PATH, "course"), os.path.join(WORKS_LIBRARY_PATH, "novel")]:
        if not os.path.exists(dir_path): os.makedirs(dir_path)
    
    token_tracker = TokenTracker(TOKEN_STATS_FILE)
    generator = ContentGenerator(TARGET_API_URL, API_HEADERS, SERVED_MODEL_IDENTIFIER, token_tracker)
    task_manager = TaskManager(generator)
    task_manager.start()
    
    print("="*50)
    print("AI Content Factory v7.0 (Web UI) Started!")
    print(f"Please open your browser to http://127.0.0.1:5000")
    print("="*50)
    app.run(host='0.0.0.0', port=5000, debug=False)